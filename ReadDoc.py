from docx import Document

# doc_path = "Hithere.docx"
# document = Document(doc_path)

# for paragraph in document.paragraphs:
#     print(paragraph.text)



def ReadDoc(filename):
    doc = Document(filename)

    paragraphs = []
    for paragraph in doc.paragraphs:
        if len(paragraph.text)>0:
            paragraphs.append(paragraph.text)
    print(paragraphs)


ReadDoc("Hithere.docx")